from __future__ import annotations

import argparse
import math
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_FORMAT_PATH = BASE_DIR / "input" / "format" / "teg-format.yaml"
TEG_OUTPUT_DIR = BASE_DIR / "output" / "test-execution-guides"


def load_yaml(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def safe_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "-", name or "")
    name = re.sub(r"\s+", " ", name).strip()
    return name[:80] or "Title"


def normalize_header(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip().lower()
    text = re.sub(r"[_\-]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text


def clean_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_lines(text: str) -> List[str]:
    if not text:
        return []
    lines = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^(?:[-•*]+\s*|\d+[\.)]\s*)", "", line).strip()
        if line:
            lines.append(line)
    return lines


def build_header_aliases(format_cfg: dict) -> Dict[str, set]:
    aliases = format_cfg["excel"]["header_aliases"]
    return {
        key: {normalize_header(x) for x in values}
        for key, values in aliases.items()
    }


def detect_header_row(ws, header_aliases: Dict[str, set], max_scan_rows: int) -> Tuple[int, Dict[str, int]]:
    best_row = 1
    best_map: Dict[str, int] = {}
    best_score = -1

    for row_idx in range(1, min(max_scan_rows, ws.max_row) + 1):
        current_map: Dict[str, int] = {}

        for col_idx in range(1, ws.max_column + 1):
            header = normalize_header(ws.cell(row=row_idx, column=col_idx).value)
            if not header:
                continue

            for logical_name, aliases in header_aliases.items():
                if header in aliases and logical_name not in current_map:
                    current_map[logical_name] = col_idx

        score = len(current_map)
        if score > best_score:
            best_score = score
            best_row = row_idx
            best_map = current_map

    if not best_map:
        raise ValueError(
            "Could not detect the header row in the Excel file. "
            "Check the sheet headers and teg-format.yaml header_aliases."
        )

    return best_row, best_map


def load_test_cases(excel_path: Path, sheet_name: Optional[str], format_cfg: dict) -> List[Dict[str, str]]:
    wb = load_workbook(excel_path, data_only=True)
    ws = wb[sheet_name] if sheet_name else wb[wb.sheetnames[0]]

    header_aliases = build_header_aliases(format_cfg)
    max_scan_rows = format_cfg["excel"].get("header_scan_rows", 10)

    header_row, header_map = detect_header_row(ws, header_aliases, max_scan_rows)

    cases: List[Dict[str, str]] = []

    for row_idx in range(header_row + 1, ws.max_row + 1):
        row_data = {
            logical_name: clean_text(ws.cell(row=row_idx, column=col_idx).value)
            for logical_name, col_idx in header_map.items()
        }

        tc_id = row_data.get("tc_id", "")
        title = row_data.get("title", "")
        steps = row_data.get("steps", "")
        expected = row_data.get("expected", "")

        if not any([tc_id, title, steps, expected]):
            continue

        if not tc_id:
            tc_id = f"TC-{len(cases) + 1}"

        if not title:
            title = f"Test Case Row {row_idx}"

        cases.append({
            "tc_id": tc_id,
            "title": title,
            "steps": steps,
            "expected": expected,
        })

    if not cases:
        raise ValueError("No test cases found in the selected sheet.")

    return cases


def apply_font(run, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_document_defaults(doc: Document, font_name: str, font_size: int) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    for style_name in ["List Bullet", "List Number"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = font_name
            style.font.size = Pt(font_size)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_document_keep_sections(doc: Document) -> None:
    body = doc._element.body
    sectPr = None

    for child in list(body):
        if child.tag.endswith("sectPr"):
            sectPr = child
        body.remove(child)

    if sectPr is not None:
        body.append(sectPr)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    font_name: str,
    font_size: int
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    apply_font(run, font_name, font_size)
    return p


def add_bullet(doc: Document, text: str, *, font_name: str, font_size: int):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    apply_font(run, font_name, font_size)
    return p


def add_number(doc: Document, text: str, *, font_name: str, font_size: int):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    apply_font(run, font_name, font_size)
    return p


def add_list_from_multiline(
    doc: Document,
    text: str,
    *,
    style: str,
    font_name: str,
    font_size: int
):
    lines = split_lines(text)
    if not lines:
        return

    for line in lines:
        if style == "number":
            add_number(doc, line, font_name=font_name, font_size=font_size)
        else:
            add_bullet(doc, line, font_name=font_name, font_size=font_size)


def build_pass_if_text(expected_text: str, prefix: str) -> str:
    lines = split_lines(expected_text)
    if not lines:
        return prefix + "Actual result matches the expected result from the test case sheet."
    return prefix + "; ".join(lines)


def add_case_block(doc: Document, case: Dict[str, str], format_cfg: dict) -> None:
    font_name = format_cfg["document"]["font_name"]
    font_size = format_cfg["document"]["font_size"]
    sections = format_cfg["sections"]

    guide_title = (
        sections["guide_title_prefix"]
        + case["tc_id"]
        + sections["guide_title_separator"]
        + case["title"]
    )

    add_paragraph(
        doc,
        guide_title,
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )

    devtools = sections["devtools_checks"]
    add_paragraph(
        doc,
        devtools["heading"],
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )
    for item in devtools["bullets"]:
        add_bullet(doc, item, font_name=font_name, font_size=font_size)

    evidence = sections["evidence_to_capture"]
    add_paragraph(
        doc,
        evidence["heading"],
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )
    for item in evidence["numbered"]:
        add_number(doc, item, font_name=font_name, font_size=font_size)
    for item in evidence["bullets"]:
        add_bullet(doc, item, font_name=font_name, font_size=font_size)

    pf = sections["pass_fail_rule"]
    add_paragraph(
        doc,
        pf["heading"],
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )
    add_bullet(
        doc,
        build_pass_if_text(case.get("expected", ""), pf["pass_if_prefix"]),
        font_name=font_name,
        font_size=font_size,
    )
    add_bullet(
        doc,
        pf["fail_if_prefix"] + pf["fail_if_text"],
        font_name=font_name,
        font_size=font_size,
    )
    add_bullet(
        doc,
        pf["blocked_if_prefix"] + pf["blocked_if_text"],
        font_name=font_name,
        font_size=font_size,
    )

    notes = sections["notes"]
    add_paragraph(
        doc,
        notes["heading"],
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )
    add_bullet(
        doc,
        notes.get("bullet", ""),
        font_name=font_name,
        font_size=font_size,
    )


def create_execution_guide(
    excel_path: Path,
    output_path: Optional[Path],
    sheet_name: Optional[str],
    format_path: Path,
    template_path: Optional[Path] = None,
) -> Path:
    format_cfg = load_yaml(format_path)

    if template_path is None:
        template_path = BASE_DIR / format_cfg["template"]["default_template"]

    if template_path.exists():
        doc = Document(template_path)
    else:
        doc = Document()

    font_name = format_cfg["document"]["font_name"]
    font_size = format_cfg["document"]["font_size"]
    blank_line_between_blocks = format_cfg["document"].get("blank_line_between_blocks", True)

    set_document_defaults(doc, font_name, font_size)
    clear_document_keep_sections(doc)

    cases = load_test_cases(excel_path, sheet_name, format_cfg)

    add_paragraph(
        doc,
        format_cfg["document"]["title"],
        bold=True,
        font_name=font_name,
        font_size=font_size,
    )

    for idx, case in enumerate(cases):
        add_case_block(doc, case, format_cfg)

        if blank_line_between_blocks and idx != len(cases) - 1:
            add_paragraph(doc, "", font_name=font_name, font_size=font_size)

    if output_path is None:
        base_title = re.sub(r"^(?:Test Case_|TC_)", "", excel_path.stem, flags=re.IGNORECASE)
        TEG_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = TEG_OUTPUT_DIR / f"Test Execution Guide_{safe_filename(base_title)}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generate a Test Execution Guide DOCX from a test case Excel file."
    )
    parser.add_argument(
        "excel_path",
        help=r"Path to the test case Excel file, e.g. D:\Project-LODI\Workflow Semi-automation\qa-test-plan\output\Test_Case_Example.xlsx",
    )
    parser.add_argument("--sheet", dest="sheet_name", help="Optional sheet name.")
    parser.add_argument("--output", dest="output_path", help="Optional output DOCX path.")
    parser.add_argument("--format", dest="format_path", default=str(DEFAULT_FORMAT_PATH), help="Path to teg-format.yaml")
    parser.add_argument("--template", dest="template_path", help="Optional DOCX template path.")
    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    excel_path = Path(args.excel_path)
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    output_path = Path(args.output_path) if args.output_path else None
    format_path = Path(args.format_path)
    template_path = Path(args.template_path) if args.template_path else None

    if not format_path.exists():
        raise FileNotFoundError(f"Format YAML not found: {format_path}")

    result = create_execution_guide(
        excel_path=excel_path,
        output_path=output_path,
        sheet_name=args.sheet_name,
        format_path=format_path,
        template_path=template_path,
    )

    print(f"✅ Generated: {result}")


if __name__ == "__main__":
    main()
