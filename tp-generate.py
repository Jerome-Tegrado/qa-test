import re
from copy import deepcopy
from pathlib import Path
from datetime import datetime

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

BASE = Path(__file__).parent
TEMPLATE = BASE / "templates" / "Test-Plan-Template.docx"
INPUT = BASE / "input" / "tp" / "tp.yaml"
TP_CUSTOMIZE_FILE = BASE / "input" / "customize" / "tp-customize.yaml"
LEGACY_TP_CUSTOMIZE_FILE = BASE / "input" / "customize" / "test-plan.yaml"
LEGACY_TP_FORMAT_FILE = BASE / "input" / "format" / "tp-format.yaml"
OUTPUT_DIR = BASE / "output"
TP_OUTPUT_DIR = OUTPUT_DIR / "test-plans"

DEFAULT_TP_CUSTOMIZATION = {
    "defaults": {
        "tp_title": "<Component / Feature Name>",
        "intro": "<1-2 short sentences on what is being tested and why. Keep it direct.>",
        "scope": {
            "in_scope": "<1-2 short sentences. State exactly what QA will validate.>",
            "out_scope": "<1 sentence. State what is not covered.>",
        },
        "objectives": ["Verify that <requirement 1>."],
        "approach_block": (
            "Methodology: Manual Testing (UI/Visual/Interaction). Add Automation Testing only if needed.\n"
            "Type of Testing: Functional Testing, UI Testing, Responsive Testing.\n"
            "Tools Used: OpenProject, Chrome DevTools. Add Playwright only if automation support is needed."
        ),
        "schedule": [
            {"phase": "Test Planning", "start": "<MM/DD or TBD>", "end": "<MM/DD or TBD>"},
            {"phase": "Test Case Design", "start": "<MM/DD or TBD>", "end": "<MM/DD or TBD>"},
            {"phase": "Test Execution", "start": "<MM/DD or TBD>", "end": "<MM/DD or TBD>"},
            {"phase": "Bug Fix Verification", "start": "<MM/DD or TBD>", "end": "<MM/DD or TBD>"},
            {"phase": "Test Completion", "start": "<MM/DD or TBD>", "end": "<MM/DD or TBD>"},
        ],
        "environment_block": (
            "Hardware/Software: Desktop / Windows 10 / Google Chrome\n"
            "Staging URL or App Version: <TBD or staging link / version>\n"
            "Test Data Sources: <Only include this line if test data source details are needed>"
        ),
        "resources": [
            {
                "role": "QA Lead",
                "name": "Sir Nico",
                "responsibilities": "<Review TP/TC, coordinate testing coverage, and perform final validation.>",
            },
            {
                "role": "QA Test Engineer",
                "name": "Jerome",
                "responsibilities": "<Design TCs, execute tests, capture evidence, and report defects.>",
            },
            {"role": "Developer", "name": "TBD", "responsibilities": "<Fix defects and support verification.>"},
        ],
        "risks": [{"risk": "<Risk 1>", "mitigation": "<Mitigation 1>"}],
        "deliverables": ["Test Plan Document", "Test Cases Document"],
        "entry": "<1-2 compact sentences describing when testing can start. Not a list.>",
        "exit": "<1-2 compact sentences describing when testing can end. Not a list.>",
    },
    "style": {
        "font": {
            "family": "Calibri",
            "sizes": {"title": 14, "heading": 12, "body": 11, "table_body": 11},
        },
        "paragraph_styles": {
            "title": {
                "alignment": "left",
                "line_spacing": 1.15,
                "space_before_pt": 0,
                "space_after_pt": 14,
            },
            "heading": {
                "alignment": "left",
                "line_spacing": 1.15,
                "space_before_pt": 18,
                "space_after_pt": 6,
            },
            "heading_table_section": {
                "alignment": "left",
                "line_spacing": 1.15,
                "space_before_pt": 18,
                "space_after_pt": 0,
            },
            "body": {
                "alignment": "justify",
                "line_spacing": 1.15,
                "space_before_pt": 0,
                "space_after_pt": 6,
            },
            "bullet": {
                "alignment": "left",
                "line_spacing": 1.15,
                "space_before_pt": 0,
                "space_after_pt": 3,
            },
            "table_cell": {
                "alignment": "left",
                "line_spacing": 1.0,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
            "multiline_block": {
                "alignment": "left",
                "line_spacing": 1.15,
                "space_before_pt": 0,
                "space_after_pt": 3,
            },
            "placeholder_clear": {
                "alignment": "left",
                "line_spacing": 1.0,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
            "empty": {
                "alignment": "left",
                "line_spacing": 1.0,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
        },
        "labels_to_bold": [
            "Purpose/Executive Summary:",
            "In Scope:",
            "Out of Scope:",
            "Methodologies:",
            "Type of Testing:",
            "Tools Used:",
            "Hardware/Software:",
            "Staging URL or App Version:",
            "Test Data Sources:",
            "Entry Criteria:",
            "Exit Criteria:",
        ],
    },
    "layout": {
        "table": {"style": "Table Grid", "width_inches": 6.5},
        "table_section_titles": [
            "5. Test Schedule",
            "7. Resources & Responsibilities",
            "8. Risks & Mitigations",
        ],
    },
    "behavior": {
        "intro_prefix_label": "Purpose/Executive Summary:",
        "ordered_blocks": {
            "environment": {
                "target_labels": [
                    "Hardware/Software:",
                    "Staging URL or App Version:",
                    "Test Data Sources:",
                ],
                "aliases": {
                    "hardwaresoftware": ["hardwareos", "hardwaresoftware"],
                    "stagingurlorappversion": ["environmenturlappversion", "stagingurl", "appversion"],
                    "testdatasources": ["testdata", "testdatasource", "testdatasources"],
                },
            },
            "approach": {
                "target_labels": ["Methodologies:", "Type of Testing:", "Tools Used:"],
                "aliases": {
                    "methodologies": ["methodology", "methodologies"],
                    "typeoftesting": ["typesoftesting", "typeoftesting"],
                    "toolsused": ["toolsused"],
                },
            },
        },
    },
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

RUNTIME_TP_CUSTOMIZATION = deepcopy(DEFAULT_TP_CUSTOMIZATION)


def safe_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "-", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name[:80]


def normalize_text(text: str) -> str:
    """
    For single-paragraph fields only.
    Converts embedded line breaks into spaces.
    """
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\n", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def load_yaml_file(file_path: Path) -> dict:
    if not file_path.exists():
        return {}
    with open(file_path, "r", encoding="utf-8") as file:
        return yaml.safe_load(file) or {}


def deep_merge(base, override):
    if isinstance(base, dict) and isinstance(override, dict):
        merged = dict(base)
        for key, value in override.items():
            if key in merged:
                merged[key] = deep_merge(merged[key], value)
            else:
                merged[key] = value
        return merged
    return override


def merge_content_defaults(defaults, override):
    if isinstance(defaults, dict):
        override_dict = override if isinstance(override, dict) else {}
        merged = {}

        for key in defaults:
            merged[key] = merge_content_defaults(defaults[key], override_dict.get(key))

        for key, value in override_dict.items():
            if key not in merged:
                merged[key] = value

        return merged

    if isinstance(defaults, list):
        if isinstance(override, list) and len(override) > 0:
            return override
        return defaults

    if override is None:
        return defaults

    if isinstance(override, str) and not override.strip():
        return defaults

    return override


def normalize_tp_customization_schema(raw_cfg: dict) -> dict:
    if not raw_cfg:
        return {}

    if "defaults" in raw_cfg or "style" in raw_cfg or "layout" in raw_cfg or "behavior" in raw_cfg:
        return raw_cfg

    # Backward compatibility: legacy tp-format.yaml was content-only top-level fields.
    return {"defaults": raw_cfg}


def load_tp_customization() -> dict:
    if TP_CUSTOMIZE_FILE.exists():
        raw_cfg = load_yaml_file(TP_CUSTOMIZE_FILE)
    elif LEGACY_TP_CUSTOMIZE_FILE.exists():
        raw_cfg = load_yaml_file(LEGACY_TP_CUSTOMIZE_FILE)
    elif LEGACY_TP_FORMAT_FILE.exists():
        raw_cfg = load_yaml_file(LEGACY_TP_FORMAT_FILE)
    else:
        raw_cfg = {}

    normalized_cfg = normalize_tp_customization_schema(raw_cfg)
    return deep_merge(DEFAULT_TP_CUSTOMIZATION, normalized_cfg)


def runtime_get(*path, default=None):
    current = RUNTIME_TP_CUSTOMIZATION
    for key in path:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current


def resolve_alignment(value: str | None, fallback=WD_ALIGN_PARAGRAPH.LEFT):
    if not value:
        return fallback
    return ALIGNMENT_MAP.get(str(value).strip().lower(), fallback)


def ensure_prefixed_label(text: str, label: str) -> str:
    normalized_text = normalize_text(text)
    if not label:
        return normalized_text

    if not normalized_text:
        return label

    if normalized_text.lower().startswith(label.lower()):
        return normalized_text

    return f"{label} {normalized_text}"


def normalize_label_key(label: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (label or "").strip().lower())


def parse_labeled_block(block_text: str) -> dict[str, str]:
    parsed: dict[str, str] = {}
    if not block_text:
        return parsed

    for raw_line in block_text.splitlines():
        line = raw_line.strip()
        if not line or ":" not in line:
            continue
        label, value = line.split(":", 1)
        key = normalize_label_key(label)
        parsed[key] = normalize_text(value)
    return parsed


def ordered_labeled_block(
    source_text: str,
    *,
    target_labels: list[str],
    aliases: dict[str, list[str]],
) -> str:
    parsed = parse_labeled_block(source_text)
    lines: list[str] = []

    for target_label in target_labels:
        target_key = normalize_label_key(target_label)
        candidate_keys = [target_key] + [normalize_label_key(x) for x in aliases.get(target_key, [])]

        value = ""
        for key in candidate_keys:
            if key in parsed and parsed[key]:
                value = parsed[key]
                break

        if value:
            lines.append(f"{target_label} {value}".rstrip())

    return "\n".join(lines)


def clean_run_text(text: str) -> str:
    """
    Cleans text already inside template runs.
    """
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\n", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def set_para_format(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing=1.15,
    space_before_pt=0,
    space_after_pt=0,
):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    paragraph.alignment = align


def set_para_format_from_style(paragraph, style_name: str):
    style_cfg = runtime_get("style", "paragraph_styles", style_name, default={}) or {}

    set_para_format(
        paragraph,
        align=resolve_alignment(style_cfg.get("alignment"), WD_ALIGN_PARAGRAPH.LEFT),
        line_spacing=style_cfg.get("line_spacing", 1.15),
        space_before_pt=style_cfg.get("space_before_pt", 0),
        space_after_pt=style_cfg.get("space_after_pt", 0),
    )


def is_heading_or_title(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    text = paragraph.text.strip()

    if not text:
        return False

    if "title" in style_name or "heading" in style_name:
        return True

    if re.match(r"^\d+(\.\d+)*\.?\s+\S+", text):
        return True

    return False


def is_bullet_paragraph(paragraph) -> bool:
    style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
    return "list bullet" in style_name or "bullet" in style_name


def clean_existing_paragraphs(doc: Document):
    for p in doc.paragraphs:
        raw = "".join(run.text for run in p.runs)
        cleaned = clean_run_text(raw)
        if p.runs:
            p.runs[0].text = cleaned
            for run in p.runs[1:]:
                run.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    raw = "".join(run.text for run in p.runs)
                    cleaned = clean_run_text(raw)
                    if p.runs:
                        p.runs[0].text = cleaned
                        for run in p.runs[1:]:
                            run.text = ""


def replace_placeholder_preserve_runs(paragraph, mapping: dict[str, str]):
    full_text = "".join(run.text for run in paragraph.runs)
    full_text = clean_run_text(full_text)
    original_text = full_text

    for key, value in mapping.items():
        if key in full_text:
            full_text = full_text.replace(key, normalize_text(value))

    if full_text == original_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(full_text)


def replace_in_doc(doc: Document, mapping: dict[str, str]):
    for p in doc.paragraphs:
        replace_placeholder_preserve_runs(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_preserve_runs(p, mapping)


def find_paragraph(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def bold_leading_label(paragraph, label_text: str):
    text = paragraph.text.strip()
    if not text.startswith(label_text):
        return

    remainder = text[len(label_text):].lstrip()

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = label_text
        paragraph.runs[0].bold = True
    else:
        r = paragraph.add_run(label_text)
        r.bold = True

    if remainder:
        r2 = paragraph.add_run(" " + remainder)
        r2.bold = False


def apply_label_bolding(doc: Document):
    labels = runtime_get("style", "labels_to_bold", default=[]) or []

    for p in doc.paragraphs:
        for label in labels:
            bold_leading_label(p, label)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for label in labels:
                        bold_leading_label(p, label)


def insert_bullets_at_placeholder(paragraph, items: list[str]):
    for item in items:
        p = paragraph.insert_paragraph_before(normalize_text(item))
        p.style = "List Bullet"
        set_para_format_from_style(p, "bullet")

    paragraph.text = ""
    set_para_format_from_style(paragraph, "placeholder_clear")


def bold_table_header_row(table):
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
            set_para_format_from_style(p, "table_cell")


def insert_table_after(paragraph, headers: list[str], rows: list[tuple]):
    paragraph.text = ""
    set_para_format_from_style(paragraph, "placeholder_clear")

    table_cfg = runtime_get("layout", "table", default={}) or {}
    parent = paragraph._parent
    table = parent.add_table(
        rows=1,
        cols=len(headers),
        width=Inches(table_cfg.get("width_inches", 6.5)),
    )
    table.style = table_cfg.get("style", "Table Grid")

    hdr_cells = table.rows[0].cells
    for i, header_value in enumerate(headers):
        hdr_cells[i].text = str(header_value)

    for row_value in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_value):
            row_cells[i].text = "" if val is None else str(val)

    paragraph._element.addnext(table._element)
    bold_table_header_row(table)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_para_format_from_style(p, "table_cell")

    return table


def insert_multiline_block_at_placeholder(
    doc: Document,
    placeholder: str,
    block_text: str,
    *,
    justified=False,
    style_name="multiline_block",
):
    """
    Replaces one placeholder paragraph with multiple paragraphs,
    one per non-empty line.
    """
    p = find_paragraph(doc, placeholder)
    if not p:
        return

    lines = [line.strip() for line in block_text.splitlines() if line.strip()]

    for line in lines:
        new_p = p.insert_paragraph_before(line)
        set_para_format_from_style(new_p, style_name)
        if justified:
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p.text = ""
    set_para_format_from_style(p, "placeholder_clear")


def format_document_spacing(doc: Document):
    """
    Desired layout:
    - Title with visible space after
    - Section heading with visible space before
    - Smaller/zero heading space-after for table sections
    """
    table_section_titles = set(runtime_get("layout", "table_section_titles", default=[]) or [])

    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name.lower() if p.style and p.style.name else ""

        if not text:
            set_para_format_from_style(p, "empty")
            continue

        if "title" in style_name or text.lower().startswith("test plan:"):
            set_para_format_from_style(p, "title")
            continue

        if is_heading_or_title(p):
            heading_style = "heading_table_section" if text in table_section_titles else "heading"
            set_para_format_from_style(p, heading_style)
            continue

        if is_bullet_paragraph(p):
            set_para_format_from_style(p, "bullet")
            continue

        set_para_format_from_style(p, "body")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_para_format_from_style(p, "table_cell")


def set_run_font(run, font_name="Calibri", font_size_pt: int | None = None):
    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)

    if run._element.rPr is None:
        run._element.get_or_add_rPr()

    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def force_document_font(doc: Document, font_name="Calibri"):
    for p in doc.paragraphs:
        for run in p.runs:
            set_run_font(run, font_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, font_name)


def apply_typography(doc: Document):
    font_family = runtime_get("style", "font", "family", default="Calibri")
    sizes = runtime_get("style", "font", "sizes", default={}) or {}

    title_size = sizes.get("title", 14)
    heading_size = sizes.get("heading", 12)
    body_size = sizes.get("body", 11)
    table_body_size = sizes.get("table_body", body_size)

    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name.lower() if p.style and p.style.name else ""

        if "title" in style_name or text.lower().startswith("test plan:"):
            size = title_size
        elif is_heading_or_title(p):
            size = heading_size
        else:
            size = body_size

        for run in p.runs:
            set_run_font(run, font_family, size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, font_family, table_body_size)


def main():
    global RUNTIME_TP_CUSTOMIZATION

    RUNTIME_TP_CUSTOMIZATION = load_tp_customization()

    if not INPUT.exists():
        raise FileNotFoundError(f"Missing input file: {INPUT}")

    input_data = load_yaml_file(INPUT)
    content_defaults = runtime_get("defaults", default={}) or {}
    data = merge_content_defaults(content_defaults, input_data)

    doc = Document(str(TEMPLATE))

    clean_existing_paragraphs(doc)

    intro_label = runtime_get("behavior", "intro_prefix_label", default="Purpose/Executive Summary:")
    intro_with_label = ensure_prefixed_label(data.get("intro", ""), intro_label)

    mapping = {
        "{{TP_TITLE}}": normalize_text(data.get("tp_title", "")),
        "{{INTRO}}": intro_with_label,
        "{{IN_SCOPE}}": normalize_text(data.get("scope", {}).get("in_scope", "")),
        "{{OUT_SCOPE}}": normalize_text(data.get("scope", {}).get("out_scope", "")),
        "{{ENTRY}}": normalize_text(data.get("entry", "")),
        "{{EXIT}}": normalize_text(data.get("exit", "")),
    }

    replace_in_doc(doc, mapping)

    environment_cfg = runtime_get("behavior", "ordered_blocks", "environment", default={}) or {}
    normalized_environment_block = ordered_labeled_block(
        data.get("environment_block", ""),
        target_labels=environment_cfg.get(
            "target_labels",
            [
                "Hardware/Software:",
                "Staging URL or App Version:",
                "Test Data Sources:",
            ],
        ),
        aliases=environment_cfg.get("aliases", {}),
    )

    insert_multiline_block_at_placeholder(
        doc,
        "{{ENVIRONMENT_BLOCK}}",
        normalized_environment_block,
        justified=False,
        style_name="multiline_block",
    )

    approach_cfg = runtime_get("behavior", "ordered_blocks", "approach", default={}) or {}
    normalized_approach_block = ordered_labeled_block(
        data.get("approach_block", ""),
        target_labels=approach_cfg.get(
            "target_labels",
            ["Methodologies:", "Type of Testing:", "Tools Used:"],
        ),
        aliases=approach_cfg.get("aliases", {}),
    )

    insert_multiline_block_at_placeholder(
        doc,
        "{{APPROACH_BLOCK}}",
        normalized_approach_block,
        justified=False,
        style_name="multiline_block",
    )

    p_obj = find_paragraph(doc, "{{OBJECTIVES}}")
    if p_obj:
        insert_bullets_at_placeholder(p_obj, data.get("objectives", []))

    p_del = find_paragraph(doc, "{{DELIVERABLES}}")
    if p_del:
        insert_bullets_at_placeholder(p_del, data.get("deliverables", []))

    p_sched = find_paragraph(doc, "{{TABLE_SCHEDULE}}")
    if p_sched:
        rows = [
            (
                normalize_text(x.get("phase", "")),
                normalize_text(x.get("start", "")),
                normalize_text(x.get("end", "")),
            )
            for x in data.get("schedule", [])
        ]
        insert_table_after(p_sched, ["Phase", "Start Date", "End Date"], rows)

    p_res = find_paragraph(doc, "{{TABLE_RESOURCES}}")
    if p_res:
        rows = [
            (
                normalize_text(x.get("role", "")),
                normalize_text(x.get("name", "")),
                normalize_text(x.get("responsibilities", "")),
            )
            for x in data.get("resources", [])
        ]
        insert_table_after(p_res, ["Role", "Name", "Responsibilities"], rows)

    p_risk = find_paragraph(doc, "{{TABLE_RISKS}}")
    if p_risk:
        rows = [
            (
                normalize_text(x.get("risk", "")),
                normalize_text(x.get("mitigation", "")),
            )
            for x in data.get("risks", [])
        ]
        insert_table_after(p_risk, ["Risk", "Mitigation"], rows)

    apply_label_bolding(doc)
    format_document_spacing(doc)

    font_family = runtime_get("style", "font", "family", default="Calibri")
    force_document_font(doc, font_family)
    apply_typography(doc)

    TP_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    title = safe_filename(data.get("tp_title", "TestPlan"))
    out_path = TP_OUTPUT_DIR / f"Test Plan_{title}.docx"
    try:
        doc.save(str(out_path))
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        out_path = TP_OUTPUT_DIR / f"Test Plan_{title}_{ts}.docx"
        doc.save(str(out_path))

    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
